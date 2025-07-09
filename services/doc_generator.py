import os
import re
from docx import Document

def generate_docx_from_template(template_string: str, mapped_data: dict, output_path: str):
    """Generate a DOCX document from template string and mapped data"""
    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    # Normalize curly quotes and braces
    template_string = re.sub(r'[""]', '"', template_string)
    template_string = re.sub(r"['']", "'", template_string)

    print("📄 Original Template:")
    print(template_string)
    print("🔑 Mapped Data:")
    for k, v in mapped_data.items():
        print(f"  {k}: {v}")

    # Replace placeholders with values
    filled_text = template_string
    for key, value in mapped_data.items():
        placeholder = f"{{{key}}}"
        filled_text = filled_text.replace(placeholder, str(value) if value else "")

    # Clean unreplaced placeholders
    unreplaced = re.findall(r"\{([^}]+)\}", filled_text)
    if unreplaced:
        print("⚠️ Unreplaced placeholders found:", unreplaced)
    filled_text = re.sub(r"\{[^}]+\}", "", filled_text)

    # Write to DOCX
    doc = Document()
    doc.styles['Normal'].paragraph_format.space_after = 0

    for line in filled_text.split("\n"):
        if not line.strip():
            continue
        p = doc.add_paragraph()
        if line.strip().isupper():
            p.add_run(line.strip()).bold = True
        else:
            p.add_run(line.strip())

    doc.save(output_path)
    print(f"\n✅ Document generated at: {output_path}")